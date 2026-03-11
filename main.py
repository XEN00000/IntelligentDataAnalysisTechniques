import pandas as pd
import docx as Document
from docx.shared import Cm 


def fit_col_width(df, table):
    table.autofit = False
    
    CHAR_TO_CM = 0.1 
    MIN_WIDTH = 1.5
    MAX_WIDTH = 8.0

    for i, col_name in enumerate(df.columns):
        max_chars = max(
            df[col_name].astype(str).map(len).max(), 
            len(str(col_name))
        )
        print(max_chars)
        
        calculated_width = max_chars * CHAR_TO_CM
        # final_width = max(MIN_WIDTH, min(calculated_width, MAX_WIDTH))
        table.columns[i].width = Cm(calculated_width)

def row_generator(df):
    yield df.columns.tolist()

    for _, row in df.iterrows():
        yield row.tolist()

def add_table_to_docx(df: pd.DataFrame, doc: Document) -> None:
    table = doc.add_table((len(df) + 1), len(df.columns))
    table.style = 'Table Grid'
    table.autofit = False

    data_for_table = row_generator(df)

    for i, row in enumerate(data_for_table):
        for j, cell in enumerate(row):
            table.cell(i, j).text = str(cell)
    
    fit_col_width(df, table)


if __name__ == "__main__":
    df = pd.read_excel("data/dataset.xlsx")
    df['Data urodzenia'] = df['Data urodzenia'].dt.strftime('%Y-%m-%d')
    df['Czy aktywny?'] = df['Czy aktywny?'].map({True: 'Tak', False: 'Nie'})

    doc = Document.Document()
    doc.add_heading("Siply document")
    add_table_to_docx(df, doc)
    doc.save("results/document.docx")