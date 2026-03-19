import pandas as pd
import json
import os
import math
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert

class DataConverter:
    def __init__(self, config_path='settings.json'):
        self.config_path = config_path
        self.settings = self.load_settings()

    def load_settings(self):
        default_settings = {
            "title": "My file",
            "font_name": "Times New Roman",
            "font_size": 9,
            "page_size": "A4",
            "landscape": True,
            "columns_order": [],
            "margins_cm": {
                "top": 1.5,
                "bottom": 1.5,
                "left": 1.5,
                "right": 1.5
            },
            "line_spacing": 1.0
        }
        if os.path.exists(self.config_path):
            with open(self.config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        else:
            self.save_settings(default_settings)
            return default_settings

    def save_settings(self, settings):
        with open(self.config_path, 'w', encoding='utf-8') as f:
            json.dump(settings, f, indent=4, ensure_ascii=False)
        self.settings = settings

    def _calculate_column_widths(self, df, available_width_cm):
        
        font_size = self.settings.get('font_size', 9)
        char_width_cm = font_size * 0.021
        padding_cm = 0.5

        min_widths = []
        text_lengths = []

        for col in df.columns:
            # sprawdzamy nagłowek zamieniając myslniki i ukośniki na spacje
            clean_header = str(col).replace('-', ' ').replace('/', ' ').replace('_', ' ')
            header_max_word = max([len(word) for word in clean_header.split()] or [1])
            
            total_chars_in_col = len(str(col))
            data_max_word = 1
            
            # sprawdzamy dane w kolumnie
            for item in df[col].astype(str):
                # neutralizujemy linki przed slice na wordy
                clean_item = item.replace('-', ' ').replace('/', ' ').replace('_', ' ')
                words = clean_item.split()
                
                if words:
                    longest_word = max([len(word) for word in words])
                    if longest_word > data_max_word:
                        data_max_word = longest_word
                        
                total_chars_in_col += len(item)
            
            # choosing minima for column
            final_max_word = max(header_max_word, data_max_word)
            min_widths.append((final_max_word * char_width_cm) + padding_cm)
            text_lengths.append(total_chars_in_col)

        # rozdawanie wolnego space-u 
        total_min_width = sum(min_widths)
        final_widths = min_widths[:]

        if total_min_width < available_width_cm:
            extra_space = available_width_cm - total_min_width
            
            # pierwiastek zapobiega mocnej faworyzowaniu jednej gigantycznej kolumny
            # weights są po to aby rozdysponowywać wolne miejsce minimalizując szkody
            # cierpieć będzie zawsze kolumna z najwięszą aktualną szerokością
            # to taki swojego rodzaju algorytm balansujący szerokości kolumns  
            weights = [math.sqrt(l) for l in text_lengths]
            total_weight = sum(weights)
            
            if total_weight > 0:
                for i in range(len(final_widths)):
                    final_widths[i] += extra_space * (weights[i] / total_weight)

        return final_widths, total_min_width

    def _set_dynamic_page_size(self, doc, required_width_cm):
        sections = doc.sections
        margins = self.settings.get('margins_cm', {"top": 1.5, "bottom": 1.5, "left": 1.5, "right": 1.5})
        
        for section in sections:
            section.top_margin = Cm(margins['top'])
            section.bottom_margin = Cm(margins['bottom'])
            section.left_margin = Cm(margins['left'])
            section.right_margin = Cm(margins['right'])
            
            section.page_width = Cm(required_width_cm)
            section.page_height = Cm(21.0) if self.settings['landscape'] else Cm(29.7)

    def convert_excel_to_docx(self, excel_path, output_docx):
        df = pd.read_excel(excel_path)
        
        if self.settings['columns_order']:
            valid_cols = [c for c in self.settings['columns_order'] if c in df.columns]
            df = df[valid_cols]

        doc = Document()
        
        margins = self.settings.get('margins_cm', {"left": 1.5, "right": 1.5})
        margins_total = margins.get('left', 1.5) + margins.get('right', 1.5)
        
        # fizyczny limit Worda - 55.8 cm
        base_page_width = 29.7 if self.settings['landscape'] else 21.0
        MAX_WORD_PAGE_WIDTH = 55.8 
        
        # calculating szerokości 
        available_width_for_table = base_page_width - margins_total
        column_widths, absolute_min_table_width = self._calculate_column_widths(df, available_width_for_table)
        
        # jesli słowa wystają za A4 to powiększamy kartkę
        target_table_width = sum(column_widths)
        final_page_width = target_table_width + margins_total

        # lock bo Word psuje dokumenty > 55.8 cm
        if final_page_width > MAX_WORD_PAGE_WIDTH:
            scale_down = (MAX_WORD_PAGE_WIDTH - margins_total) / target_table_width
            column_widths = [w * scale_down for w in column_widths]
            final_page_width = MAX_WORD_PAGE_WIDTH

        # ustawienia kartki
        self._set_dynamic_page_size(doc, final_page_width)

        if self.settings['title']:
            title = doc.add_heading(self.settings['title'], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.name = self.settings['font_name']
                run.font.color.rgb = None
                run.font.bold = True

        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        table.autofit = False

        # Wymuszanie szerokości na poziomie obiektu kolumny 
        # bo wrod lubi sobie sam poustawiać aka rozwalić te szerokości
        for i, width in enumerate(column_widths):
            table.columns[i].width = Cm(width)

        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].text = str(column_name)
            paragraph = hdr_cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True

        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                cell = row_cells[i]
                cell.text = str(value) if pd.notna(value) else ""
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # wymuszamy wyliczone szerokości na każdą komórkę z osobna
        # bo były przypadki że każda komórka miała osbie inną szerokość w jednej kolumnie XD
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(column_widths[i])
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = self.settings['line_spacing']
                    for run in paragraph.runs:
                        run.font.name = self.settings['font_name']
                        run.font.size = Pt(self.settings['font_size'])

        self._add_page_numbers(doc)
        doc.save(output_docx)
        print(f"Pomyślnie utworzono dokument Word (Szerokość kartki: {final_page_width:.2f} cm): {output_docx}")

    def _add_page_numbers(self, doc):
        for section in doc.sections:
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = paragraph.add_run("Strona ")
            
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = " PAGE "
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            t = OxmlElement('w:t')
            t.text = "1"
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run._r.append(t)
            run._r.append(fldChar3)

    def convert_docx_to_pdf(self, input_docx, output_pdf):
        try:
            convert(input_docx, output_pdf)
            print(f"Pomyślnie przekonwertowano do PDF: {output_pdf}")
        except Exception as e:
            print(f"Błąd podczas konwersji do PDF: {e}")
